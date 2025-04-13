import pymupdf
import pytesseract
import docx
from PIL import Image
import io
from bs4 import BeautifulSoup as bs
import epub
import torch
from PIL import Image
import time
import textract
from transformers import ViTFeatureExtractor, AutoTokenizer, VisionEncoderDecoderModel
import pathlib
import os

_device = "cuda" if torch.cuda.is_available() else "cpu"


#-------------------Texts--------------------------------------------------
def extract_pdf_text(file_path):
    start_time = time.time()
    try:
        # Attempt to open the PDF file
        doc = pymupdf.open(file_path)
        text_extracted = ""

        # Loop through all pages
        for page_num in range(len(doc)):
            try:
                page = doc.load_page(page_num)
                text = page.get_text("text")
                text_extracted += text

                # If we have enough text, stop
                if len(text_extracted) >= 1200:
                    return text_extracted
            except Exception as e:
                print(f"Error extracting text from page {page_num}: {e}")
                continue
        return text_extracted

    except Exception as e:
        return None


def extract_docx_text(file_path):
    doc = docx.Document(file_path)
    text = ""
    start_time = time.time()
    # Extract text from paragraphs first
    for para in doc.paragraphs:
        text += para.text + "\n"
        if len(text) >= 1200:
            break

    # If the text is still under 500 characters, add text from tables
    if len(text) < 1200:
        table_text = ""
        for table in doc.tables:
            for row in table.rows:
                row_data = [cell.text.strip() for cell in row.cells]
                table_text += " | ".join(row_data) + "\n"
                if len(text + table_text) >= 1200:
                    break
            if len(text + table_text) >= 1200:
                break

        # Add table text to the final output, ensuring total length is ≤ 500
        text += table_text

    return text


def extract_txt_text(file_path):
    start_time = time.time()
    try:
        with open(file_path, "r", encoding="utf-8") as file:
            content = file.read()
        return content[:1200]
    except UnicodeDecodeError:
        print(f"Error decoding file {file_path}. Trying with a different encoding.")
        with open(file_path, "r", encoding="latin1") as file:
            content = file.read()
        return content[:1200]


def extract_doc_text(file_path):
    start_time = time.time()
    text = textract.process(file_path).decode("utf-8")
    return text



def extract_tex_text(file_path):
    start_time = time.time()
    with open(file_path, "r", encoding="utf-8") as file:

        return file.read()[:1200]

def extract_epub_text(file_path):
    book = epub.read_epub(file_path)
    text = []
    char_count = 0
    start_time = time.time()
    for item in book.get_items():
        if item.get_type() == epub.EpubItem.TEXT:
            soup = bs(item.get_body_content(), 'html.parser')
            extracted_text = soup.get_text()
            text.append(extracted_text)
            char_count += len(extracted_text)
            if char_count >= 1200:
                break

    return "".join(text)[:1200]

#-------------------Images--------------------------------------------------



# Global configs for maximum speed
os.environ["TOKENIZERS_PARALLELISM"] = "false"  # Avoid deadlocks
if torch.cuda.is_available():
    torch.backends.cudnn.benchmark = True
    torch.backends.cudnn.deterministic = False
    torch.backends.cuda.matmul.allow_tf32 = True  # Allow TF32 on Ampere GPUs
    torch.backends.cudnn.allow_tf32 = True

# Use lightweight model
MODEL_NAME = "nlpconnect/vit-gpt2-image-captioning"

# Initialize global variables
device = torch.device("cuda" if torch.cuda.is_available() else "cpu")
tokenizer = None
feature_extractor = None
model = None
is_initialized = False

def init_model():
    """Initialize and optimize the model only once"""
    global tokenizer, feature_extractor, model, is_initialized

    if is_initialized:
        return

    print("Loading model...")
    # Load fastest versions of tokenizer and feature_extractor
    tokenizer = AutoTokenizer.from_pretrained(MODEL_NAME)
    feature_extractor = ViTFeatureExtractor.from_pretrained(MODEL_NAME)

    # Load and optimize model
    model = VisionEncoderDecoderModel.from_pretrained(MODEL_NAME)
    model.to(device)
    model.eval()

    # Apply aggressive speedups
    if device.type == "cuda":
        model = model.half()  # Use FP16 precision

    # Quantize the model (this reduces model size and improves speed)
    model = torch.quantization.quantize_dynamic(
        model, {torch.nn.Linear}, dtype=torch.qint8
    )

    # Disable all gradients
    for param in model.parameters():
        param.requires_grad = False

    # JIT/compile optimizations for newer PyTorch
    if hasattr(torch, 'compile'):
        try:
            model = torch.compile(model, mode="reduce-overhead", fullgraph=True)
        except Exception as e:
            print(f"Torch compile error: {e}")

    is_initialized = True
    print("Model ready")

def preprocess_image(image_path, target_size=(224, 224)):
    """Preprocess image with maximum efficiency"""
    # Handle both filepath and PIL image
    if isinstance(image_path, (str, pathlib.PosixPath)):
        image_path = str(image_path)  # Convert to string if it's a PosixPath
        image = Image.open(image_path).convert('RGB')
    else:
        image = image_path

    if image.size != target_size:
        image = image.resize(target_size, Image.BILINEAR)

    # Fast feature extraction
    inputs = feature_extractor(images=image, return_tensors="pt")

    # Create attention mask (1 for real tokens, 0 for padding)
    attention_mask = inputs['pixel_values'].new_ones(inputs['pixel_values'].shape[:2])

    # Move to device and optimize precision
    if device.type == "cuda":
        inputs = {k: v.to(device).half() for k, v in inputs.items()}
        attention_mask = attention_mask.to(device).half()  # Ensure it matches the device
    else:
        inputs = {k: v.to(device) for k, v in inputs.items()}
        attention_mask = attention_mask.to(device)

    # Add the attention mask to the inputs
    inputs['attention_mask'] = attention_mask

    return inputs

def generate_caption(inputs):
    """Generate caption with minimal settings for speed"""
    # Ensure no_grad context for inference
    with torch.no_grad(), torch.cuda.amp.autocast() if device.type == "cuda" else nullcontext():
        output_ids = model.generate(
            inputs["pixel_values"],         # Minimal length for speed
            num_beams=1,             # Greedy search (fastest)
            do_sample=False,
            early_stopping=True,
            use_cache=True,
            return_dict_in_generate=False,
            output_scores=False
        )

    # Fast decoding
    caption = tokenizer.decode(output_ids[0], skip_special_tokens=True)
    return caption

def caption_image(image_path):
    """Main function to caption an image with timing"""
    # Make sure model is initialized
    if not is_initialized:
        init_model()

    # Time the actual processing
    start_time = time.time()

    # Process pipeline
    start_time2 = time.time()

    inputs = preprocess_image(image_path)
    print(f"Preprocessing Time: {time.time() - start_time2:.4f} seconds")
    start_time1 = time.time()
    caption = generate_caption(inputs)
    print(f"Generation Time: {time.time() - start_time1:.4f} seconds")

    # Calculate timing
    inference_time = time.time() - start_time
    return caption, inference_time

# Nullcontext for PyTorch compatibility
class nullcontext:
    def __enter__(self): return None
    def __exit__(self, *args): pass