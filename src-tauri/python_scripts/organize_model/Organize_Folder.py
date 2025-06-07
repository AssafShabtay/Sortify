import torch
from pathlib import Path
import time
from deep_translator import GoogleTranslator
import umap
import hdbscan
import pandas as pd
import concurrent.futures
import os
from sentence_transformers import SentenceTransformer
import nltk
from nltk.corpus import stopwords
from nltk.stem import WordNetLemmatizer
import re
import optuna
import numpy as np
from hdbscan.validity import validity_index
import pymupdf
import docx
from bs4 import BeautifulSoup as bs
import epub
import textract
import sys
import codecs
import json
import random
from PIL import Image
from io import BytesIO
import cairosvg
import pytesseract
from sklearn.feature_extraction.text import TfidfVectorizer, CountVectorizer
from collections import Counter
import string
from sklearn.metrics.pairwise import cosine_similarity


np.random.seed(42)
random.seed(42)
torch.manual_seed(42)
torch.cuda.manual_seed_all(42)
os.environ['PYTHONHASHSEED'] = '42'

torch.backends.cudnn.deterministic = True
torch.backends.cudnn.benchmark = False


folder_path = sys.argv[1]
output_json_path = sys.argv[2]
toplevel_folders_as_one = sys.argv[3]

error_output_path = Path(output_json_path).parent / "errors.json"
unsupported_extensions_output_path = Path(output_json_path).parent / "unsupported_extensions.json"

error_collection = {}

nltk.download('stopwords')
nltk.download('wordnet')

translator = GoogleTranslator(source='auto', target='en')
model = SentenceTransformer("sentence-transformers/all-mpnet-base-v2")

unsupported_extensions = []

num_cores = os.cpu_count()

max_workers_suggestion = min(32, num_cores * 2 + 1)

treat_toplevel_folders_as_one = True if toplevel_folders_as_one == "true" else False

# make the prints be in utf-8
if sys.platform == 'win32':
    sys.stdout = codecs.getwriter('utf-8')(sys.stdout.buffer, 'strict')
    sys.stderr = codecs.getwriter('utf-8')(sys.stderr.buffer, 'strict')

def add_error(error_message, context=""):
    """
    Add an error to the collection, tracking both unique errors and their count
    
    Args:
        error_message: The error message string
        context: Additional context about where the error occurred
    """
    # Create a key that combines error type and message for uniqueness
    if isinstance(error_message, Exception):
        error_type = type(error_message).__name__
        error_msg = str(error_message)
    else:
        error_type = "Error"
        error_msg = str(error_message)
    
    error_key = f"{error_type}: {error_msg}"
    
    # Add or update the error in the collection
    if error_key in error_collection:
        error_collection[error_key]["count"] += 1
        # Only add new contexts
        if context and context not in error_collection[error_key]["contexts"]:
            error_collection[error_key]["contexts"].append(context)
    else:
        error_collection[error_key] = {
            "count": 1,
            "contexts": [context] if context else []
        }

def print_error_summary():
    """Print a summary of all collected errors"""
    print("\n=== ERROR SUMMARY ===")
    print(f"Total unique errors: {len(error_collection)}")
    
    for error_key, error_data in error_collection.items():
        print(f"\n{error_key}")
        print(f"Occurred {error_data['count']} times")
        if error_data["contexts"]:
            print("Sample contexts:")
            for i, context in enumerate(error_data["contexts"][:3]):  # Show up to 3 contexts
                print(f"  - {context}")
            if len(error_data["contexts"]) > 3:
                print(f"  - ... and {len(error_data['contexts']) - 3} more locations")
    print("======================\n")

#-------------------Texts--------------------------------------------------
def extract_pdf_text(file_path):
    try:
        doc = pymupdf.open(file_path)
        text_extracted = ""

        for page_num in range(len(doc)):
            try:
                page = doc.load_page(page_num)
                text = page.get_text("text")
                text_extracted += text

                if len(text_extracted) >= 1200:
                    return text_extracted
            except Exception as e:
                add_error(e, f"PDF page extraction: {file_path}, page {page_num}")
                continue
        return text_extracted

    except Exception as e:
        add_error(e, f"PDF file opening: {file_path}")
        return None

def extract_doc_text(file_path):
    try:
        text = textract.process(file_path).decode("utf-8")
        return text
    except Exception as e:
        add_error(e, f"DOC extraction: {file_path}")
        return None

def extract_docx_text(file_path):
    try:
        doc = docx.Document(file_path)
        text = ""
     
        for para in doc.paragraphs:
            text += para.text + "\n"
            if len(text) >= 1200:
                break

        if len(text) < 1200:
            table_text = ""
            for table in doc.tables:
                for row in table.rows:
                    row_data = [cell.text.strip() for cell in row.cells]
                    table_text += " | ".join(row_data) + "\n"
                    if len(text + table_text) >= 1200:
                        break
                if len(text + table_text) >= 1200:
                    break

            text += table_text

        return text
    except Exception as e:
        add_error(e, f"DOCX extraction: {file_path}")
        return None

def extract_txt_text(file_path):
    try:
        with open(file_path, "r", encoding="utf-8") as file:
            content = file.read()
        return content[:1200]
    except UnicodeDecodeError as e:
        try:
            add_error(e, f"UTF-8 decoding: {file_path}")
            print(f"Error decoding file {file_path}. Trying with a different encoding.")
            with open(file_path, "r", encoding="latin1") as file:
                content = file.read()
            return content[:1200]
        except Exception as e2:
            add_error(e2, f"Latin1 decoding: {file_path}")
            return None
    except Exception as e:
        add_error(e, f"TXT file reading: {file_path}")
        return None

def extract_tex_text(file_path):
    try:
        with open(file_path, "r", encoding="utf-8") as file:
            return file.read()[:1200]
    except Exception as e:
        add_error(e, f"TEX file reading: {file_path}")
        return None

def extract_epub_text(file_path):
    try:
        book = epub.read_epub(file_path)
        text = []
        char_count = 0
        for item in book.get_items():
            if item.get_type() == epub.EpubItem.TEXT:
                soup = bs(item.get_body_content(), 'html.parser')
                extracted_text = soup.get_text()
                text.append(extracted_text)
                char_count += len(extracted_text)
                if char_count >= 1200:
                    break

        return "".join(text)[:1200]
    except Exception as e:
        add_error(e, f"EPUB extraction: {file_path}")
        return None

#-------------------Images--------------------------------------------------
def caption_image(image_path):
    try:
        extension = os.path.splitext(image_path)[-1].lower()
        file_name = os.path.splitext(os.path.basename(image_path))[0]
        if extension == "svg":
            png_data = BytesIO()
            cairosvg.svg2png(url=str(image_path), write_to=png_data)
            png_data.seek(0)
            image_input = Image.open(png_data)
        else:
            image_input = Image.open(image_path)

        text = pytesseract.image_to_string(image_input)
        text = text + "   " + file_name 
        return text
    except Exception as e:
        add_error(e, f"Image captioning: {image_path}")
        return None

def extract_file_summary(file_path):
    print(f"Processing file: {file_path}")
    start_time = time.time()
    extension = os.path.splitext(file_path)[-1].lower()
    try:
        if extension == ".pdf":
            text = extract_pdf_text(Path(file_path))
        elif extension == ".docx":
            text = extract_docx_text(Path(file_path))
        elif extension == ".txt":
            text = extract_txt_text(Path(file_path))
        elif extension == ".doc":
            text = extract_doc_text(Path(file_path))
        elif extension == ".tex":
            text = extract_tex_text(Path(file_path))
        elif extension == ".epub":
            text = extract_epub_text(Path(file_path))
        elif extension in (".jpg", ".jpeg", ".png", ".gif", ".bmp", ".tiff", ".tif", ".webp", ".ico", ".heif", ".heic", ".avif", ".eps", ".dds", ".dis", ".im", ".mpo", ".msp", ".pxc", ".pfm", ".ppm", ".tga", ".spider", ".sgi", ".xbm", "psd", ".svg"):
            text = caption_image(Path(file_path))
            print(f"Image caption: {text}")
            return text
        else:
            unsupported_extensions.append(extension)
            return None

        if text:
            try:
                translated_text = translator.translate(text[:500])
                if translated_text:
                    return translated_text
                else:
                    error_msg = f"Translation failed for: {file_path}"
                    add_error(error_msg, "Translation")
                    print(error_msg)
                    return None
            except Exception as e:
                add_error(e, f"Translation: {file_path}")
                return None
        else:
            error_msg = f"No text extracted from: {file_path}"
            add_error(error_msg, "Text extraction")
            print(error_msg)
            return None

    except Exception as e:
        add_error(e, f"General file processing: {file_path}")
        print(f"Error processing {file_path}: {e}")
        return None
    finally:
        print(f"Processing time for {file_path}: {time.time() - start_time:.2f} seconds")

def process_directory(folder_path, max_workers=4):
    summaries = []
    print(f"Processing folder: {folder_path}")
    
    file_paths = []
    for root, dirs, files in os.walk(folder_path):
        for file in files:
            file_path = os.path.join(root, file)
            if os.path.isfile(file_path):
                file_paths.append(file_path)

    with concurrent.futures.ThreadPoolExecutor(max_workers) as executor:
        future_to_file = {executor.submit(extract_file_summary, file_path): file_path for file_path in file_paths}
        for future in concurrent.futures.as_completed(future_to_file):
            file_path = future_to_file[future]
            try:
                summary_temp = future.result()
                if summary_temp:
                    summaries.append(summary_temp[:500])
            except Exception as e:
                add_error(e, f"Directory processing: {file_path}")
                print(f"Error processing {file_path}: {e}")

    if not summaries:
        return "No summaries were generated for this folder."
    combined_summary = " | ".join(filter(None, summaries))

    return combined_summary
    
def run_model_organizer(folder_path, dict_as_one, max_workers=4):
    file_summaries = {}
    MAX_FILE_SIZE = 50 * 1024 * 1024  # 50MB
    start_time = time.time()

    file_paths = []
    top_level_dir_paths = []
    if dict_as_one:
        with os.scandir(folder_path) as entries:
          for entry in entries:
              entry_path = entry.path

              if entry.is_file() and os.path.getsize(entry_path) <= MAX_FILE_SIZE:
                  file_paths.append(entry_path)
              elif entry.is_dir():
                  top_level_dir_paths.append(entry_path)
    else:
      print(f"Processing top-level files in: {folder_path}")
      for root, dirs, files in os.walk(folder_path):
         for file in files:
            file_path = os.path.join(root, file)
            if os.path.isfile(file_path) and os.path.getsize(file_path) <= MAX_FILE_SIZE:
              file_paths.append(file_path)

    print(f"Processing {len(file_paths)} top-level files")
    with concurrent.futures.ThreadPoolExecutor(max_workers=max_workers) as executor:
        future_to_file = {executor.submit(extract_file_summary, file_path): file_path for file_path in file_paths}

        for future in concurrent.futures.as_completed(future_to_file):
            file_path = future_to_file[future]
            try:
                summary = future.result()
                if summary:
                    file_summaries[file_path] = summary
            except Exception as e:
                add_error(e, f"Top-level file processing: {file_path}")
                print(f"Error processing top-level file {file_path}: {e}")
                
    if dict_as_one:
      print(f"Processing {len(top_level_dir_paths)} directories")
      with concurrent.futures.ThreadPoolExecutor(max_workers=max(1, max_workers//2)) as executor:
          future_to_dir = {executor.submit(process_directory, dir_path, max_workers=2): dir_path for dir_path in top_level_dir_paths}

          for future in concurrent.futures.as_completed(future_to_dir):
              dir_path = future_to_dir[future]
              try:
                  summary = future.result()
                  file_summaries[dir_path] = summary
              except Exception as e:
                  add_error(e, f"Directory execution: {dir_path}")
                  print(f"Error processing directory {dir_path}: {e}")

    print(f"Total processing time: {time.time() - start_time:.2f} seconds")

    return file_summaries

#-------------------Improved Cluster Naming System--------------------------------------------------

def extract_key_phrases(texts, max_phrases=5):
    """Extract key phrases using n-grams and frequency analysis"""
    
    try:
        vectorizer = CountVectorizer(
            ngram_range=(2, 3),
            max_features=100,
            stop_words='english'
        )
        
        all_text = " ".join(texts)

        X = vectorizer.fit_transform([all_text])
        
        feature_names = vectorizer.get_feature_names_out()
        counts = X.toarray()[0]
        
        phrase_counts = [(feature_names[i], counts[i]) for i in range(len(counts)) if counts[i] > 0]
        phrase_counts.sort(key=lambda x: x[1], reverse=True)
        
        return [phrase for phrase, count in phrase_counts[:max_phrases]]
    except Exception as e:
        add_error(e, "Key phrase extraction")
        return []

def find_semantic_themes(cluster_df, embeddings, stop_words):
    """Find semantic themes using embeddings and clustering within the cluster"""
    try:
        if len(cluster_df) < 3:
            return []
        
        cluster_indices = cluster_df.index.tolist()
        cluster_embeddings = embeddings[cluster_indices]
        
        if len(cluster_embeddings) > 20:
            cluster_embeddings = cluster_embeddings[:20]
            cluster_indices = cluster_indices[:20]
        
        centroid = np.mean(cluster_embeddings, axis=0)
        
        similarities = cosine_similarity([centroid], cluster_embeddings)[0]
        top_indices = np.argsort(similarities)[-3:][::-1]
        
        representative_texts = cluster_df.iloc[top_indices]['texts'].tolist()
        
        important_terms = []
        for text in representative_texts:
            words = text.lower().split()
            for word in words:
                if (len(word) > 4 and 
                    word not in stop_words and 
                    word.isalpha()):
                    important_terms.append(word)
        
        term_counts = Counter(important_terms)
        return [term for term, count in term_counts.most_common(5)]
        
    except Exception as e:
        add_error(e, "Semantic theme extraction")
        return []
    
def analyze_file_extensions(cluster_df):
    """Get file type from extension"""
    try:
        extensions = []
        for path in cluster_df['path']:
            ext = os.path.splitext(path)[1].lower()
            if ext:
                extensions.append(ext)
        
        if not extensions:
            return None
        
        ext_counts = Counter(extensions)
        most_common = ext_counts.most_common(3)
        
        ext_mapping = {
            '.pdf': 'Documents',
            '.docx': 'Documents',
            '.doc': 'Documents',
            '.txt': 'Text Files',
            '.tex': 'LaTeX',
            '.epub': 'E-books',
            '.jpg': 'Images',
            '.jpeg': 'Images',
            '.png': 'Images',
            '.svg': 'Vector Graphics'
        }
        
        content_types = []
        for ext, count in most_common:
            if ext in ext_mapping:
                content_types.append(ext_mapping[ext])
        
        if content_types:
            return Counter(content_types).most_common(1)[0][0]
        
        return None
        
    except Exception as e:
        add_error(e, "File extension analysis")
        return None

def extract_directory_patterns(cluster_df):
    """Extract common directory patterns"""
    try:
        dirs = []
        for path in cluster_df['path']:
            dir_path = os.path.dirname(path)
            if dir_path:
                # Get last two directory components
                parts = Path(dir_path).parts[-2:]
                dirs.extend(parts)
        
        if not dirs:
            return None
        
        # Find common directory names
        dir_counts = Counter(dirs)
        common_dirs = [d for d, count in dir_counts.most_common(3) 
                      if count >= len(cluster_df) * 0.3]  # At least 30% of files
        
        if common_dirs:
            # Clean
            cleaned = [d.replace('_', ' ').replace('-', ' ').title() 
                      for d in common_dirs[:2]]
            return " - ".join(cleaned)
        
        return None
        
    except Exception as e:
        add_error(e, "Directory pattern extraction")
        return None

def clean_cluster_name(name):
    """Clean and format cluster name"""
  
    name = name.lstrip()
    
    unwanted_chars = '#%&{}\\<>*?/$!\'"":@+`|='
    
    for char in unwanted_chars:
        name = name.replace(char, ' ')
    
    name = re.sub(r'[^\w\s\-().,]', ' ', name, flags=re.ASCII)
    
    name = ' '.join(name.split())
    
    words = name.split()
    seen = set()
    cleaned_words = []
    for word in words:
        word_lower = word.lower()
        if word_lower not in seen:
            seen.add(word_lower)
            cleaned_words.append(word)
    
    name = ' '.join(cleaned_words)
    
    if len(name) > 60:
        name = name[:57] + "..."
    
    name = string.capwords(name)
    
    return name

def generate_descriptive_name(cluster_id, cluster_df, embeddings, stop_words):

    name_candidates = []
    
    # Extractions ranked by priority, Higher number = Higher priority

    dir_pattern = extract_directory_patterns(cluster_df)
    if dir_pattern: 
        name_candidates.append((3, dir_pattern))
    

    content_type = analyze_file_extensions(cluster_df)
    if content_type:
        name_candidates.append((1, content_type))
    
    themes = find_semantic_themes(cluster_df, embeddings, stop_words)
    if themes:
        theme_name = " ".join(themes[:3]).title()
        name_candidates.append((2, theme_name))
    
    texts = cluster_df['cleaned_texts'].tolist()
    key_phrases = extract_key_phrases(texts, max_phrases=3)
    if key_phrases:
        phrase_name = " - ".join(key_phrases[:2]).title()
        name_candidates.append((2, phrase_name))
    
    try:
        tfidf = TfidfVectorizer(
            max_df=0.8,
            min_df=2,
            max_features=50,
            ngram_range=(1, 2),
            sublinear_tf=True
        )
        
        tfidf_matrix = tfidf.fit_transform(texts)
        
        avg_scores = np.mean(tfidf_matrix.toarray(), axis=0)
        top_indices = np.argsort(avg_scores)[-5:][::-1]
        feature_names = tfidf.get_feature_names_out()
        
        tfidf_terms = [feature_names[i] for i in top_indices]
        if tfidf_terms:
            tfidf_name = " ".join(tfidf_terms[:3]).title()
            name_candidates.append((1, tfidf_name))
            
    except Exception as e:
        add_error(e, f"Advanced TF-IDF for cluster {cluster_id}")
    

    if name_candidates:
        name_candidates.sort(key=lambda x: x[0], reverse=True)

        best_name = name_candidates[0][1]
        
        if len(name_candidates) > 1 and name_candidates[1][0] >= 2:
            secondary = name_candidates[1][1]
            if secondary not in best_name:
                best_name = f"{best_name} ({secondary})"

        if content_type and content_type not in best_name:
            best_name = f"{content_type}: {best_name}"
        
        best_name = clean_cluster_name(best_name)
        
        return best_name
    
    # Fallback to generic name
    return f"Cluster {cluster_id}"

def name_all_clusters(dataset, embeddings):
    cluster_names = {}
    used_names = set()
    
    cluster_names[-1] = "Uncategorized"
    unique_clusters = sorted(dataset['label'].unique())
    
    for cluster_id in unique_clusters:
        if cluster_id == -1:
            continue
            
        cluster_df = dataset[dataset['label'] == cluster_id]
        
        if len(cluster_df) < 2:
            cluster_names[cluster_id] = f"Small Group {cluster_id}"
            continue
        
        base_name = generate_descriptive_name(cluster_id, cluster_df, embeddings, stop_words)
        
        # Handle duplicates
        final_name = base_name
        counter = 1
        while final_name in used_names:
            counter += 1
            final_name = f"{base_name} ({counter})"
        
        used_names.add(final_name)
        cluster_names[cluster_id] = final_name
        
        print(f"Cluster {cluster_id}: '{final_name}' ({len(cluster_df)} files)")
    
    return cluster_names

try:
    file_summaries = run_model_organizer(folder_path, treat_toplevel_folders_as_one)
    corrupt_files = []
    dataset = pd.DataFrame(
          [(path, data) for path, data in file_summaries.items()],
          columns=['path', 'texts']
    )

    dataset['filename'] = dataset['path'].apply(lambda x: Path(x).stem)
    dataset = dataset.sort_values(by='filename').reset_index(drop=True)

    stop_words = set(stopwords.words('english'))
    lemmatizer = WordNetLemmatizer()

    def clean_text(text):
        if not isinstance(text, str):
            return "" 
        text = text.lower()
        text = re.sub(r'[^a-z0-9\s]', '', text, flags=re.A)
        words = text.split()
        cleaned_lemmatized_words = []
        for word in words:
            if word and word not in stop_words: 
                lemmatized_word = lemmatizer.lemmatize(word)
                cleaned_lemmatized_words.append(lemmatized_word)

        return ' '.join(cleaned_lemmatized_words)
        
    
    
    none_mask = dataset['texts'].isnull()
    corrupt_files.extend(dataset[none_mask]['path'].tolist())

    dataset['cleaned_texts'] = dataset['texts'].apply(clean_text)

    none_mask = dataset['cleaned_texts'].str.len() <= 3
    corrupt_files.extend(dataset[none_mask]['path'].tolist())
    
    dataset = dataset[dataset['cleaned_texts'].str.len() > 3]
    cleaned_texts = dataset['cleaned_texts'].tolist()

    corrupt_files = list(set(corrupt_files))
    del none_mask

    X = model.encode(cleaned_texts, convert_to_numpy=True)
    if len(X) < 6:
        error_msg = "Not enough files"
        add_error(error_msg, "Input validation")
        sys.stderr.write(error_msg)
        
        # Save error collection before exiting
        with open(error_output_path, "w", encoding="utf-8") as f:
            json.dump(error_collection, f, ensure_ascii=False, indent=2)

        sys.exit(1)

    num_components = min(len(X) - 3, 15)
    umap_model = umap.UMAP(n_components=num_components, random_state=42)
    X_reduced = umap_model.fit_transform(X)

    X_reduced = X_reduced.astype(np.float64)
    n_samples = X_reduced.shape[0]

    # Find best parameters
    def objective(trial):
        try:
            min_samples = trial.suggest_int("min_samples", 2, 100 if n_samples >= 50 else n_samples - 1)
            min_cluster_size = trial.suggest_int("min_cluster_size", 2, min(n_samples - 1, 15))
            alpha = trial.suggest_float("alpha", 0.5, 1.5)
            epsilon = trial.suggest_float("cluster_selection_epsilon", 0.5, 1.0)
            metric = trial.suggest_categorical("metric", ["euclidean", "manhattan", "l2"])
            p = trial.suggest_float("p", 1.0, 3.0)

            clusterer = hdbscan.HDBSCAN(
                min_samples=min_samples,
                min_cluster_size=min_cluster_size,
                alpha=alpha,
                cluster_selection_epsilon=epsilon,
                metric=metric,
                p=p,
                cluster_selection_method="eom",
                prediction_data=True,
                allow_single_cluster=False
            )

            labels = clusterer.fit_predict(X_reduced)
            
            if len(set(labels)) <= 1:
                return -1  

            return validity_index(X_reduced, labels)
        except Exception as e:
            add_error(e, f"Optuna trial {trial.number}")
            return -1  # Return a default value to continue optimization

    sampler = optuna.samplers.NSGAIISampler(seed=42) 
    study = optuna.create_study(direction="maximize", sampler=sampler)
    
    try:
        study.optimize(objective, n_trials=400)
        
        trials = [t.number for t in study.trials]
        scores = [t.value for t in study.trials]

        print(f"Best DBCV Score: {study.best_value}")
        print(f"Best Parameters: {study.best_params}")

        best_params = study.best_params

        best_clusterer = hdbscan.HDBSCAN(
            min_samples=best_params["min_samples"],
            min_cluster_size=best_params["min_cluster_size"],
            alpha=best_params["alpha"],
            cluster_selection_epsilon=best_params["cluster_selection_epsilon"],
            metric=best_params["metric"],
            p=best_params["p"],
            cluster_selection_method="eom",
            prediction_data=True,
            allow_single_cluster=False
        )

        labels = best_clusterer.fit_predict(X_reduced)

        dataset['label'] = labels
        print("Cluster labels assigned to DataFrame.")

        unique_labels = sorted(dataset['label'].unique())
        n_clusters_found = len(unique_labels)
        if -1 in unique_labels:
            n_clusters_found -= 1

        if n_clusters_found < 1:
            print("HDBSCAN found no meaningful clusters (only noise or a single cluster). Exiting.")
            output_data = {
                "cluster_assignments": [{"path": row['path'], "label": int(row['label'])} for index, row in dataset.iterrows()],
                "cluster_names": {str(l): "No meaningful clusters found" for l in unique_labels}
            }
            if unique_labels == [-1]:
                 output_data["cluster_names"] = {"-1": "No clusters found"}
            elif len(unique_labels) == 1 and unique_labels[0] != -1:
                 output_data["cluster_names"] = {str(unique_labels[0]): "Single Cluster Found"}

            with open(output_json_path, "w", encoding="utf-8") as f:
                json.dump(output_data, f, ensure_ascii=False, indent=2)
            with open(error_output_path, "w", encoding="utf-8") as f:
                json.dump(error_collection, f, ensure_ascii=False, indent=2)
            
            sys.exit(0)

        print(f"Clustering complete. Found {len(unique_labels)} clusters (including noise if present).")

        #-------------------Use Improved Cluster Naming--------------------------------------------------

        print("Starting improved cluster naming process...")
    
        # Generate names for all clusters
        cluster_names_map = name_all_clusters(dataset, X)
        
        print("Cluster naming complete.")

        #-------------------Output JSON -------------------------------------------------- 

        unique_unsupported_extensions = list(set(unsupported_extensions))
        
        with open(unsupported_extensions_output_path, "w", encoding="utf-8") as f:
            json.dump(unsupported_extensions, f, ensure_ascii=False, indent=2)

        cluster_assignments = [
            {
                "path": row['path'], 
                "cluster_name": cluster_names_map[row['label']]
            }
            for index, row in dataset.iterrows()
        ]

        for path in corrupt_files:
            cluster_assignments.append({
                "path": path,
                "cluster_name": "Corrupt"
            })

        # Combine into the final output structure
        output_data = cluster_assignments

        try:
            with open(output_json_path, "w", encoding="utf-8") as f:
                json.dump(output_data, f, ensure_ascii=False, indent=2)
        except Exception as e:
            add_error(e, f"Writing output JSON to {output_json_path}")
            print(f"Error writing output JSON to {output_json_path}: {e}")
            sys.exit(1)

    except Exception as e:
        add_error(e, "Optimization process")
        print(f"Error in optimization process: {e}")

except Exception as e:
    add_error(e, "Main execution")
    print(f"Error in main execution: {e}")

print(f"Writing error summary to {error_output_path}")
try:
    with open(error_output_path, "w", encoding="utf-8") as f:
        json.dump(error_collection, f, ensure_ascii=False, indent=2)
except Exception as e:
    print(f"Error writing error collection to JSON file: {e}")

print("Processing complete. JSON output generated.")
#print_error_summary()
sys.exit(0)